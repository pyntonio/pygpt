import openai
from docx import Document
import os

class PortfolioGenerator:
    def __init__(self, api_key):
        openai.api_key = api_key

    def generate_portfolio(self, budget, goals, timeline, preferences=""):
        prompt = (
            f"Create a detailed investment portfolio based on the following parameters:\n"
            f"- Budget: {budget}\n- Goals: {goals}\n- Timeline: {timeline}\n"
            f"- Preferences: {preferences}\n"
            f"Include asset allocation, risk analysis, and potential returns."
        )
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.7
        )
        return response['choices'][0]['message']['content']

    def generate_word_report(self, portfolio_text, file_name="portfolio_report.docx"):
        document = Document()
        document.add_heading("Investment Portfolio", level=1)
        document.add_paragraph(portfolio_text)
        file_path = os.path.join("reports", file_name)
        document.save(file_path)
        return file_path
