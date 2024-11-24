import openai
from docx import Document
from fastapi.responses import FileResponse
import os

class StockAnalysis:
    def __init__(self, api_key):
        openai.api_key = api_key

    def analyze_stock_or_index(self, stock_name, period, additional_context=""):
        prompt = (
            f"Perform a detailed analysis of the stock or index '{stock_name}' for the period '{period}'. "
            f"Include historical performance, key financial metrics, and future projections. {additional_context}"
        )
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.7
        )
        return response['choices'][0]['message']['content']

    def generate_word_report(self, content: str, file_path: str) -> str:
        # Crea la directory se non esiste
        directory = os.path.dirname(file_path)
        if directory and not os.path.exists(directory):
            os.makedirs(directory)
        
        # Crea il documento Word
        document = Document()
        document.add_heading("Stock or Index Analysis", level=1)
        document.add_paragraph(content)
        
        # Salva il file
        document.save(file_path)
        return file_path

